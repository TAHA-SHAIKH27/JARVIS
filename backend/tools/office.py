import asyncio
import os
from typing import Any, Dict, List
from docx import Document
from system_ops import WORK_DIR


class Office:
    """Word/DOCX generation tools."""

    @staticmethod
    async def create_docx(content: str, title: str = "Document", headings: List[str] = None,
                          lists: List[str] = None, tables: List[List] = None,
                          save_path: str = "") -> Dict[str, Any]:
        """Create a Microsoft Word .docx file."""
        try:
            # Determine save path
            if not save_path:
                # Default to work_files
                filename = f"{title.replace(' ', '_')}_{os.path.basename(os.getcwd())}.docx"
                save_path = os.path.join(WORK_DIR, filename)
            elif not save_path.endswith('.docx'):
                save_path = save_path if save_path.endswith('.docx') else save_path + '.docx'

            # Ensure directory exists
            os.makedirs(os.path.dirname(save_path), exist_ok=True)

            # Create document
            doc = Document()

            # Add title
            if title:
                doc.add_heading(title, level=0)

            # Add headings
            if headings:
                for heading in headings:
                    doc.add_heading(heading, level=1)

            # Add content paragraphs
            for paragraph in content.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)

            # Add lists
            if lists:
                for item in lists:
                    doc.add_paragraph(item, style='List Bullet')

            # Add tables
            if tables:
                for table_data in tables:
                    if isinstance(table_data, list) and table_data:
                        rows = len(table_data)
                        cols = max(len(row) for row in table_data) if table_data else 0
                        table = doc.add_table(rows=rows, cols=cols)
                        for i, row_data in enumerate(table_data):
                            for j, cell_text in enumerate(row_data):
                                if j < cols:
                                    table.rows[i].cells[j].text = str(cell_text)

            doc.save(save_path)
            return {"status": "success", "message": f"Word document created: {save_path}", "path": save_path}
        except Exception as e:
            return {"status": "error", "message": f"Failed to create Word document: {str(e)}"}

    @staticmethod
    async def open_document(path: str) -> Dict[str, Any]:
        """Open a document (attempt to launch with default application)."""
        try:
            if os.path.exists(path):
                os.startfile(path) if os.name == 'nt' else None
                return {"status": "success", "message": f"Opened document: {path}"}
            return {"status": "error", "message": f"Document not found: {path}"}
        except Exception as e:
            return {"status": "error", "message": f"Failed to open document: {str(e)}"}

    @staticmethod
    async def verify_document(path: str) -> Dict[str, Any]:
        """Verify a document exists and has content."""
        try:
            if os.path.exists(path):
                doc = Document(path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                return {"status": "success", "message": f"Document verified: {path}", "paragraphs": len(paragraphs)}
            return {"status": "error", "message": f"Document not found: {path}"}
        except Exception as e:
            return {"status": "error", "message": f"Failed to verify document: {str(e)}"}